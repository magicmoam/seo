"""Client-facing SEO pitch report generator (Word .docx)."""

from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from src.models import (
    AgentResponse,
    SEOStrategy,
    WebsiteAnalysisResult,
)

# ---------------------------------------------------------------------------
# Brand constants
# ---------------------------------------------------------------------------

_ACCENT = RGBColor(0xD4, 0xB8, 0x95)  # warm gold
_DARK = RGBColor(0x1A, 0x1D, 0x2E)
_GREEN = RGBColor(0x2E, 0x7D, 0x32)
_YELLOW = RGBColor(0xF5, 0x9E, 0x0B)
_RED = RGBColor(0xC6, 0x28, 0x28)
_GRAY = RGBColor(0x61, 0x61, 0x61)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_LIGHT_BG = RGBColor(0xF5, 0xF5, 0xF5)


def _score_color(score: int) -> RGBColor:
    if score >= 70:
        return _GREEN
    if score >= 40:
        return _YELLOW
    return _RED


def _score_label(score: int) -> str:
    if score >= 80:
        return "Excellent"
    if score >= 60:
        return "Good"
    if score >= 40:
        return "Needs Work"
    return "Critical"


def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    shading.append(shd)


def _add_styled_paragraph(doc, text: str, *, size: int = 11, bold: bool = False,
                          color: RGBColor | None = None, space_after: int = 6,
                          alignment=None, space_before: int = 0) -> None:
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def _style_table(table):
    """Apply consistent styling to a table."""
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if row_idx == 0:
                _set_cell_shading(cell, "1A1D2E")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = _WHITE
                        run.bold = True
                        run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _add_cover(doc: Document, domain: str):
    for _ in range(6):
        doc.add_paragraph()

    _add_styled_paragraph(doc, "SEO INTELLIGENCE REPORT", size=28, bold=True,
                          color=_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # Accent line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 40)
    run.font.color.rgb = _ACCENT
    run.font.size = Pt(10)

    _add_styled_paragraph(doc, "", size=14, space_after=20)

    _add_styled_paragraph(doc, domain.upper(), size=20, bold=True,
                          color=_ACCENT, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    date_str = datetime.now(timezone.utc).strftime("%B %d, %Y")
    _add_styled_paragraph(doc, f"Prepared {date_str}", size=12,
                          color=_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _add_styled_paragraph(doc, "trySEO.ai SEO Intelligence", size=12,
                          color=_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _add_styled_paragraph(doc, "Confidential", size=10,
                          color=_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Sections for SEOStrategy
# ---------------------------------------------------------------------------

def _add_executive_summary_strategy(doc: Document, r: SEOStrategy):
    doc.add_heading("Executive Summary", level=1)

    # Score box
    score = r.overall_health_score
    color = _score_color(score)

    table = doc.add_table(rows=1, cols=3)
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(6)
    table.columns[2].width = Cm(5)

    c0 = table.cell(0, 0)
    c0.text = ""
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(score))
    run.font.size = Pt(48)
    run.font.color.rgb = color
    run.bold = True
    run2 = p.add_run("/100")
    run2.font.size = Pt(18)
    run2.font.color.rgb = _GRAY

    c1 = table.cell(0, 1)
    c1.text = ""
    p1 = c1.paragraphs[0]
    run = p1.add_run(f"Overall Health: {_score_label(score)}")
    run.font.size = Pt(14)
    run.bold = True
    p2 = c1.add_paragraph()
    run2 = p2.add_run(f"Timeline: {r.expected_timeline_to_results}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = _GRAY

    c2 = table.cell(0, 2)
    c2.text = ""
    p = c2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"${r.total_monthly_investment_usd:,.0f}/mo")
    run.font.size = Pt(18)
    run.font.color.rgb = _ACCENT
    run.bold = True
    p2 = c2.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Recommended Investment")
    run2.font.size = Pt(9)
    run2.font.color.rgb = _GRAY

    doc.add_paragraph()
    _add_styled_paragraph(doc, r.executive_summary, size=11, color=_DARK, space_after=12)

    _add_styled_paragraph(doc, f"ROI Projection: {r.roi_projection}", size=10,
                          color=_GRAY, space_after=16)

    doc.add_page_break()


def _add_priority_actions(doc: Document, r: SEOStrategy):
    doc.add_heading("Priority Action Plan", level=1)
    _add_styled_paragraph(
        doc,
        "The following actions are ranked by impact and feasibility. "
        "Quick wins are highlighted to show immediate value.",
        size=10, color=_GRAY, space_after=12,
    )

    # Group by timeline
    timelines = {}
    for a in r.priority_actions:
        tl = a.timeline.replace("_", " ").title()
        timelines.setdefault(tl, []).append(a)

    for timeline, actions in timelines.items():
        doc.add_heading(timeline, level=2)

        table = doc.add_table(rows=1 + len(actions), cols=5)
        table.columns[0].width = Cm(5)
        table.columns[1].width = Cm(2)
        table.columns[2].width = Cm(4)
        table.columns[3].width = Cm(3)
        table.columns[4].width = Cm(2)

        headers = ["Action", "Priority", "Expected Impact", "Source", "Est. Cost"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h

        for idx, a in enumerate(actions, 1):
            table.cell(idx, 0).text = a.action
            table.cell(idx, 1).text = a.priority.upper()
            table.cell(idx, 2).text = a.expected_impact
            table.cell(idx, 3).text = a.agent.replace("_", " ").title()
            table.cell(idx, 4).text = f"${a.estimated_cost_usd:,.0f}"

            # Color-code priority
            pri_cell = table.cell(idx, 1)
            for p in pri_cell.paragraphs:
                for run in p.runs:
                    if a.priority == "critical":
                        run.font.color.rgb = _RED
                    elif a.priority == "high":
                        run.font.color.rgb = _YELLOW

        _style_table(table)
        doc.add_paragraph()

    doc.add_page_break()


def _add_technical_audit(doc: Document, r: SEOStrategy):
    doc.add_heading("Technical SEO Assessment", level=1)
    audit = r.technical_audit

    _add_styled_paragraph(
        doc,
        f"Mobile Readiness: {audit.mobile_readiness.title()} | "
        f"Estimated Fix Effort: {audit.estimated_total_effort}",
        size=10, color=_GRAY, space_after=12,
    )

    # Top priority actions
    doc.add_heading("Priority Technical Fixes", level=2)
    for i, action in enumerate(audit.priority_actions[:5], 1):
        _add_styled_paragraph(doc, f"{i}. {action}", size=10, color=_DARK, space_after=4)

    doc.add_paragraph()

    # Fixes table
    if audit.fixes:
        doc.add_heading("Detailed Findings", level=2)
        table = doc.add_table(rows=1 + len(audit.fixes), cols=5)
        headers = ["Category", "Issue", "Impact", "Recommendation", "Effort"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h

        for idx, f in enumerate(audit.fixes, 1):
            table.cell(idx, 0).text = f.category.replace("_", " ").title()
            table.cell(idx, 1).text = f.issue
            table.cell(idx, 2).text = f.impact.upper()
            table.cell(idx, 3).text = f.recommendation
            table.cell(idx, 4).text = f.estimated_effort

            # Color impact
            imp_cell = table.cell(idx, 2)
            for p in imp_cell.paragraphs:
                for run in p.runs:
                    if f.impact == "critical":
                        run.font.color.rgb = _RED
                    elif f.impact == "high":
                        run.font.color.rgb = _YELLOW

        _style_table(table)

    # Schema recommendations
    if audit.schema_recommendations:
        doc.add_paragraph()
        doc.add_heading("Schema Markup Recommendations", level=2)
        for s in audit.schema_recommendations:
            _add_styled_paragraph(doc, f"{s.schema_type}: {s.description}",
                                  size=10, color=_DARK, space_after=6)

    doc.add_page_break()


def _add_content_strategy(doc: Document, r: SEOStrategy):
    doc.add_heading("Content Strategy", level=1)
    authority = r.topical_authority

    _add_styled_paragraph(
        doc,
        f"Niche: {authority.niche} | "
        f"Total Pages Recommended: {authority.estimated_total_pages}",
        size=10, color=_GRAY, space_after=12,
    )

    _add_styled_paragraph(doc, authority.summary, size=11, color=_DARK, space_after=12)

    # Clusters overview
    doc.add_heading("Topical Authority Clusters", level=2)
    for i, cluster in enumerate(authority.clusters, 1):
        doc.add_heading(f"Cluster {i}: {cluster.cluster_name}", level=3)
        _add_styled_paragraph(
            doc,
            f"Pillar: {cluster.pillar_page}\n"
            f"Target Keyword: {cluster.pillar_keyword} "
            f"(Volume: {cluster.pillar_search_volume})",
            size=10, color=_DARK, space_after=6,
        )
        _add_styled_paragraph(
            doc,
            f"Supporting Pages: {len(cluster.supporting_pages)}",
            size=10, color=_GRAY, space_after=4,
        )
        for page in cluster.supporting_pages[:5]:
            _add_styled_paragraph(doc, f"  - {page}", size=9, color=_GRAY, space_after=2)
        if len(cluster.supporting_pages) > 5:
            _add_styled_paragraph(
                doc,
                f"  ... and {len(cluster.supporting_pages) - 5} more",
                size=9, color=_GRAY, space_after=4,
            )

    # Content calendar highlight
    doc.add_paragraph()
    doc.add_heading("12-Week Publishing Calendar", level=2)
    cal = r.content_calendar
    _add_styled_paragraph(
        doc,
        f"Publishing Frequency: {cal.publishing_frequency} | "
        f"Total Investment: ${cal.total_estimated_cost_usd:,.0f}",
        size=10, color=_GRAY, space_after=8,
    )

    # Show first 12 entries in a table
    entries = cal.entries[:12]
    if entries:
        table = doc.add_table(rows=1 + len(entries), cols=6)
        headers = ["Week", "Title", "Type", "Priority", "Words", "Cost"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h

        for idx, e in enumerate(entries, 1):
            table.cell(idx, 0).text = str(e.week)
            table.cell(idx, 1).text = e.title[:50] + ("..." if len(e.title) > 50 else "")
            table.cell(idx, 2).text = e.content_type.title()
            table.cell(idx, 3).text = e.priority.title()
            table.cell(idx, 4).text = f"{e.estimated_word_count:,}"
            table.cell(idx, 5).text = f"${e.estimated_cost_usd:,.0f}"

        _style_table(table)

    doc.add_page_break()


def _add_backlink_strategy(doc: Document, r: SEOStrategy):
    doc.add_heading("Link Building Strategy", level=1)
    bl = r.backlink_strategy

    _add_styled_paragraph(
        doc,
        f"Monthly Link Target: {bl.monthly_link_target} links | "
        f"Estimated Monthly Cost: ${bl.estimated_monthly_cost_usd:,.0f}",
        size=10, color=_GRAY, space_after=12,
    )

    _add_styled_paragraph(doc, bl.summary, size=11, color=_DARK, space_after=12)

    # Opportunities table
    if bl.opportunities:
        doc.add_heading("Link Acquisition Opportunities", level=2)
        table = doc.add_table(rows=1 + len(bl.opportunities), cols=5)
        headers = ["Strategy", "Target", "Difficulty", "DA Range", "Cost"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h

        for idx, o in enumerate(bl.opportunities, 1):
            table.cell(idx, 0).text = o.strategy.replace("_", " ").title()
            table.cell(idx, 1).text = o.target_description[:60]
            table.cell(idx, 2).text = o.difficulty.title()
            table.cell(idx, 3).text = o.estimated_domain_authority_range
            table.cell(idx, 4).text = f"${o.estimated_cost_usd:,.0f}"

        _style_table(table)

    # Linkable assets
    if bl.linkable_assets:
        doc.add_paragraph()
        doc.add_heading("Recommended Linkable Assets", level=2)
        for asset in bl.linkable_assets:
            _add_styled_paragraph(
                doc,
                f"{asset.asset_type.replace('_', ' ').title()}: {asset.title}",
                size=10, bold=True, color=_DARK, space_after=2,
            )
            _add_styled_paragraph(
                doc,
                f"{asset.description} (Link potential: {asset.estimated_links_potential}, "
                f"Cost: ${asset.estimated_production_cost_usd:,.0f})",
                size=9, color=_GRAY, space_after=8,
            )

    doc.add_page_break()


def _add_investment(doc: Document, r: SEOStrategy):
    doc.add_heading("Investment & Returns", level=1)

    # Investment summary
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)

    c0 = table.cell(0, 0)
    c0.text = ""
    _set_cell_shading(c0, "FFF8E1")
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"${r.total_monthly_investment_usd:,.0f}")
    run.font.size = Pt(24)
    run.font.color.rgb = _ACCENT
    run.bold = True
    p2 = c0.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Monthly Investment")
    run2.font.size = Pt(10)
    run2.font.color.rgb = _GRAY

    c1 = table.cell(0, 1)
    c1.text = ""
    _set_cell_shading(c1, "E8F5E9")
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"${r.total_setup_cost_usd:,.0f}")
    run.font.size = Pt(24)
    run.font.color.rgb = _GREEN
    run.bold = True
    p2 = c1.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("One-Time Setup Cost")
    run2.font.size = Pt(10)
    run2.font.color.rgb = _GRAY

    doc.add_paragraph()

    # Cost breakdown table
    if r.cost_breakdown:
        doc.add_heading("Cost Breakdown", level=2)
        table = doc.add_table(rows=1 + len(r.cost_breakdown), cols=5)
        headers = ["Category", "Description", "Monthly", "One-Time", "Notes"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h

        for idx, c in enumerate(r.cost_breakdown, 1):
            table.cell(idx, 0).text = c.category.replace("_", " ").title()
            table.cell(idx, 1).text = c.description
            table.cell(idx, 2).text = f"${c.monthly_cost_usd:,.0f}"
            table.cell(idx, 3).text = f"${c.one_time_cost_usd:,.0f}"
            table.cell(idx, 4).text = c.notes

        _style_table(table)

    doc.add_paragraph()

    # ROI & Timeline
    doc.add_heading("Expected Returns", level=2)
    _add_styled_paragraph(doc, f"Timeline to Results: {r.expected_timeline_to_results}",
                          size=11, bold=True, color=_DARK, space_after=6)
    _add_styled_paragraph(doc, f"ROI Projection: {r.roi_projection}",
                          size=11, color=_DARK, space_after=12)

    # Risk factors
    if r.risk_factors:
        doc.add_heading("Risk Factors", level=2)
        for risk in r.risk_factors:
            _add_styled_paragraph(doc, f"- {risk}", size=10, color=_GRAY, space_after=3)

    doc.add_page_break()


def _add_path_to_80_section(doc: Document, p80):
    """Add path-to-80 section if available."""
    doc.add_heading("Score Improvement Roadmap", level=1)
    _add_styled_paragraph(
        doc,
        f"Current Score: {p80.current_score} | Target: {p80.target_score} | "
        f"Projected: {p80.projected_score}",
        size=11, bold=True, color=_DARK, space_after=8,
    )
    _add_styled_paragraph(doc, p80.quick_wins_summary, size=10, color=_GRAY, space_after=12)

    if p80.steps:
        table = doc.add_table(rows=1 + len(p80.steps), cols=5)
        headers = ["#", "Action", "Category", "Points", "Effort"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h

        running = p80.current_score
        for idx, s in enumerate(p80.steps, 1):
            running += s.estimated_points
            table.cell(idx, 0).text = str(idx)
            table.cell(idx, 1).text = s.action
            table.cell(idx, 2).text = s.category.title()
            table.cell(idx, 3).text = f"+{s.estimated_points}"
            table.cell(idx, 4).text = s.effort.replace("_", " ").title()

            # Green the points column
            pts_cell = table.cell(idx, 3)
            for p in pts_cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = _GREEN

        _style_table(table)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Sections for WebsiteAnalysisResult (site audit pitch)
# ---------------------------------------------------------------------------

def _add_executive_summary_audit(doc: Document, r: WebsiteAnalysisResult):
    doc.add_heading("Executive Summary", level=1)

    score = r.overall_score
    color = _score_color(score)

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(6)
    table.columns[1].width = Cm(10)

    c0 = table.cell(0, 0)
    c0.text = ""
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(score))
    run.font.size = Pt(56)
    run.font.color.rgb = color
    run.bold = True
    run2 = p.add_run("/100")
    run2.font.size = Pt(20)
    run2.font.color.rgb = _GRAY
    p2 = c0.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p2.add_run(_score_label(score))
    run3.font.size = Pt(14)
    run3.font.color.rgb = color
    run3.bold = True

    c1 = table.cell(0, 1)
    c1.text = ""
    _add_cell_paragraph(c1, r.page_title, size=14, bold=True, color=_DARK)
    _add_cell_paragraph(c1, r.url, size=9, color=_GRAY)
    _add_cell_paragraph(c1, "", size=6)
    _add_cell_paragraph(c1, r.summary, size=10, color=_DARK)
    _add_cell_paragraph(
        c1,
        f"{r.word_count:,} words | {r.internal_links} internal links | "
        f"{r.external_links} external links",
        size=9, color=_GRAY,
    )

    doc.add_paragraph()

    # Category scores row
    cats = [
        ("Performance", r.performance_score),
        ("SEO", r.seo_score),
        ("Content", r.content_score),
        ("Technical", r.technical_score),
    ]
    cat_table = doc.add_table(rows=2, cols=4)
    for i, (name, label) in enumerate(cats):
        cat_table.cell(0, i).text = name
        cat_table.cell(1, i).text = label.title()
        for p in cat_table.cell(0, i).paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = _GRAY
                run.font.name = "Calibri"
        for p in cat_table.cell(1, i).paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(12)
                run.bold = True
                run.font.name = "Calibri"
                score_map = {"excellent": _GREEN, "good": _GREEN, "fair": _YELLOW, "poor": _RED}
                run.font.color.rgb = score_map.get(label.lower(), _GRAY)

    # Rubric numeric scores if available
    if r.rubric:
        doc.add_paragraph()
        rubric_table = doc.add_table(rows=2, cols=4)
        for i, cat_bd in enumerate(r.rubric.categories[:4]):
            rubric_table.cell(0, i).text = cat_bd.category.title()
            rubric_table.cell(1, i).text = f"{cat_bd.score}/100 ({cat_bd.passed_count}/{cat_bd.total_count})"
            for p in rubric_table.cell(0, i).paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = _GRAY
                    run.font.name = "Calibri"
            for p in rubric_table.cell(1, i).paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.bold = True
                    run.font.name = "Calibri"
                    run.font.color.rgb = _score_color(cat_bd.score)

    doc.add_page_break()


def _add_rubric_breakdown(doc: Document, r: WebsiteAnalysisResult):
    """Add per-criterion rubric breakdown."""
    if not r.rubric:
        return

    doc.add_heading("Detailed Scoring Breakdown", level=1)
    _add_styled_paragraph(
        doc,
        f"Each criterion is evaluated against industry-standard benchmarks. "
        f"Rubric v{r.rubric.rubric_version} | {r.rubric.overall_score}/100 overall.",
        size=10, color=_GRAY, space_after=12,
    )

    for cat_bd in r.rubric.categories:
        color = _score_color(cat_bd.score)
        doc.add_heading(
            f"{cat_bd.category.title()} - {cat_bd.score}/100 "
            f"({cat_bd.passed_count}/{cat_bd.total_count} passed)",
            level=2,
        )

        table = doc.add_table(rows=1 + len(cat_bd.criteria), cols=4)
        table.columns[0].width = Cm(1)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(6)
        table.columns[3].width = Cm(5)

        headers = ["", "Criterion", "Finding", "Recommendation"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h

        for idx, cr in enumerate(cat_bd.criteria, 1):
            status = "PASS" if cr.passed else "FAIL"
            table.cell(idx, 0).text = status
            table.cell(idx, 1).text = f"{cr.criterion_name} ({cr.score})"
            table.cell(idx, 2).text = cr.finding[:120] + ("..." if len(cr.finding) > 120 else "")
            table.cell(idx, 3).text = cr.recommendation[:120] if cr.recommendation else ""

            # Color status
            status_cell = table.cell(idx, 0)
            for p in status_cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = _GREEN if cr.passed else _RED
                    run.bold = True

        _style_table(table)
        doc.add_paragraph()

    doc.add_page_break()


def _add_issues_section(doc: Document, r: WebsiteAnalysisResult):
    """Add issues table for site audit."""
    if not r.issues:
        return

    doc.add_heading("Issues Identified", level=1)
    _add_styled_paragraph(
        doc,
        "Each issue below directly impacts your search visibility and organic traffic potential.",
        size=10, color=_GRAY, space_after=12,
    )

    # Count by severity
    critical = sum(1 for i in r.issues if i.severity == "critical")
    warning = sum(1 for i in r.issues if i.severity == "warning")
    info = sum(1 for i in r.issues if i.severity == "info")

    if critical:
        _add_styled_paragraph(doc, f"{critical} critical issue(s) require immediate attention.",
                              size=11, bold=True, color=_RED, space_after=6)

    table = doc.add_table(rows=1 + len(r.issues), cols=4)
    headers = ["Severity", "Issue", "Description", "Recommendation"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h

    for idx, iss in enumerate(r.issues, 1):
        table.cell(idx, 0).text = iss.severity.upper()
        table.cell(idx, 1).text = iss.issue
        table.cell(idx, 2).text = iss.description[:100] + ("..." if len(iss.description) > 100 else "")
        table.cell(idx, 3).text = iss.recommendation[:100] + ("..." if len(iss.recommendation) > 100 else "")

        sev_cell = table.cell(idx, 0)
        for p in sev_cell.paragraphs:
            for run in p.runs:
                if iss.severity == "critical":
                    run.font.color.rgb = _RED
                elif iss.severity == "warning":
                    run.font.color.rgb = _YELLOW
                run.bold = True

    _style_table(table)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Next steps / CTA
# ---------------------------------------------------------------------------

def _add_next_steps(doc: Document, domain: str):
    doc.add_heading("Next Steps", level=1)

    _add_styled_paragraph(
        doc,
        "This report identifies specific, actionable improvements that will directly "
        "impact your search rankings, organic traffic, and revenue.",
        size=11, color=_DARK, space_after=16,
    )

    steps = [
        ("1. Review Priority Actions",
         "Focus on quick wins first - these deliver the fastest ROI with minimal effort."),
        ("2. Technical Foundation",
         "Address critical technical issues before investing in content. "
         "A strong technical foundation ensures search engines can properly crawl and index your site."),
        ("3. Content Investment",
         "Follow the recommended content calendar to build topical authority. "
         "Consistent, high-quality content is the highest-ROI SEO investment."),
        ("4. Link Building",
         "Implement the link acquisition strategy to build domain authority. "
         "Quality backlinks remain a top-3 ranking factor."),
        ("5. Ongoing Monitoring",
         "Track progress with monthly audits. SEO is iterative - "
         "continuous optimization compounds results over time."),
    ]

    for title, desc in steps:
        _add_styled_paragraph(doc, title, size=12, bold=True, color=_DARK, space_after=2)
        _add_styled_paragraph(doc, desc, size=10, color=_GRAY, space_after=12)

    doc.add_paragraph()
    _add_styled_paragraph(
        doc,
        "Ready to implement? Contact us to get started.",
        size=14, bold=True, color=_ACCENT, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8,
    )
    _add_styled_paragraph(
        doc, "tryseo.ai",
        size=12, color=_ACCENT, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
    )

    # Footer
    doc.add_paragraph()
    _add_styled_paragraph(
        doc,
        "This report was generated by trySEO.ai SEO Intelligence. "
        "Scores and recommendations are based on automated analysis of publicly available web data "
        "and industry-standard benchmarks. Results should be validated with domain-specific context.",
        size=8, color=_GRAY, space_after=0,
    )


def _add_cell_paragraph(cell, text: str, *, size: int = 11, bold: bool = False,
                        color: RGBColor | None = None):
    """Add a styled paragraph to a table cell."""
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_client_report(response: AgentResponse) -> bytes:
    """Generate a client-facing Word document report. Returns .docx bytes."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Style headings
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        heading_style.font.color.rgb = _DARK
        if level == 1:
            heading_style.font.size = Pt(22)
        elif level == 2:
            heading_style.font.size = Pt(16)
        else:
            heading_style.font.size = Pt(13)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    result = response.result
    domain = response.query

    if isinstance(result, SEOStrategy):
        domain = result.domain
        _add_cover(doc, domain)
        _add_executive_summary_strategy(doc, result)
        _add_priority_actions(doc, result)
        _add_technical_audit(doc, result)
        _add_content_strategy(doc, result)
        _add_backlink_strategy(doc, result)
        _add_investment(doc, result)
        if result.path_to_80:
            _add_path_to_80_section(doc, result.path_to_80)
        _add_next_steps(doc, domain)

    elif isinstance(result, WebsiteAnalysisResult):
        domain = result.url
        _add_cover(doc, domain)
        _add_executive_summary_audit(doc, result)
        _add_rubric_breakdown(doc, result)
        _add_issues_section(doc, result)
        if result.path_to_80:
            _add_path_to_80_section(doc, result.path_to_80)
        _add_next_steps(doc, domain)

    else:
        # Fallback: just dump the data
        _add_cover(doc, domain)
        doc.add_heading("Analysis Results", level=1)
        _add_styled_paragraph(doc, response.result.model_dump_json(indent=2),
                              size=9, color=_GRAY, space_after=0)
        _add_next_steps(doc, domain)

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
